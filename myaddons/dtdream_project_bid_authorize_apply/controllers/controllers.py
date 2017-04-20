# -*- coding: utf-8 -*-

from openerp.http import request,serialize_exception

from openerp.addons.web.controllers.main import ExcelExport
import zipfile
import os
from cStringIO import StringIO
from openerp import http
from openerp.http import request
try:
    import cPickle as pickle
except ImportError:
    import pickle
import base64
class Mymodule(ExcelExport):
    @http.route('/mymodule/mymodule/', auth='public')
    def index(self,**kw):
        fruits = http.request.env['dtdream_project_bid_authorize_apply.fruits']
        return http.request.render("dtdream_project_bid_authorize_apply.index",
                {'fruits': fruits.search([])})

    @http.route('/dtdream_authorization/dtdream_shouquan_export', type='http', auth='public')
    def dtdream_auth_export(self,data,token):
        active_ids = data.split(',')
        dtdream_authorization_apply = request.env['dtdream.project.bid.authorize.apply'].search([('id', '=', int(active_ids[0]))])#获取当前项目授权记录
        from docx import Document
        import sys
        reload(sys)
        sys.setdefaultencoding('utf-8')
        auth = dtdream_authorization_apply
        auth.attachment_ids
        project_name = auth.bidding_rep_pro_name
        bidding_number = auth.bidding_number
        date = auth.bidding_time.split(u'-')
        replace_content = ['招标单位名称','项目名称','bidnum','2016年6月12日','year','month','day','渠道']
        replace_content_dict = {'招标单位名称':auth.bidding_company,'渠道':[authorization.content for authorization in auth.authorization ],'项目名称':project_name,'bidnum':bidding_number,'year':date[0],'month':date[1],'day':date[2]}
        if replace_content_dict['渠道']:
            z = zipfile.ZipFile('authorizations.zip', 'w')#新建或清空一个zip文件用于写入需要压缩的文件
            for authorization in replace_content_dict['渠道']:#替paragraph中换字符串
                doc = Document('myaddons/dtdream_project_bid_authorize_apply/static/file/shouquan.docx')#载入docx模板
                paragraphs = doc.paragraphs
                for p in paragraphs:
                    for run in p.runs:
                        for rc in replace_content:
                            if rc == u'渠道' and u'渠道' in run.text and authorization:
                                pass
                                run.text = run.text.replace(rc,authorization)

                            elif rc in run.text and rc and replace_content_dict[rc]:
                                run.text=run.text.replace(rc,replace_content_dict[rc])
                for t in doc.tables:#替换table中的字符串
                    for row in t.rows:
                        for cell in row.cells:
                            for rc in replace_content:
                                if rc == '渠道' and '渠道' in cell.text:
                                    cell.text = cell.text.replace(rc,authorization)
                                elif rc in cell.text:
                                    cell.text=cell.text.replace(rc,replace_content_dict[rc])
                doc.save(u'{0}_{1}.docx'.format(project_name,authorization))#暂存docx文件
                z.write(u'{0}_{1}.docx'.format(project_name,authorization))#将docx文件写入zip文件
                os.remove(u'{0}_{1}.docx'.format(project_name,authorization))# 删除文件
            z.close()
            f= open('authorizations.zip')
            data = f.read()#读取压缩文件的数据
            f.close()
            return request.make_response(
                [data, data],
                headers=[
                    ('Content-Disposition', 'attachment; filename="%s"'
                     % 'authorization.zip'),
                    ('Content-Type', self.content_type)
                ],
                cookies={'fileToken': token}
            )
        else:
            doc = Document('myaddons/dtdream_project_bid_authorize_apply/static/file/shouquan.docx')
            paragraphs = doc.paragraphs
            for p in paragraphs:
                for run in p.runs:
                    for rc in replace_content:
                        if rc == '渠道' and '渠道' in run.text:
                            run.text = run.text.replace(rc,'')
                        elif rc in run.text :
                            run.text=run.text.replace(rc,replace_content_dict[rc])
            for t in doc.tables:
                for row in t.rows:
                    for cell in row.cells:
                        for rc in replace_content:
                            if rc == '渠道' and '渠道' in cell.text:
                                cell.text = cell.text.replace(rc,'')
                            elif rc in cell.text:
                                cell.text=cell.text.replace(rc,replace_content_dict[rc])
            z = zipfile.ZipFile('authorizations.zip', 'w')  # 新建或清空一个zip文件用于写入需要压缩的文件
            doc.save(u'{0}_{1}.docx'.format(project_name, replace_content_dict['渠道'][0]))  # 暂存docx文件
            z.write(u'{0}_{1}.docx'.format(project_name, replace_content_dict['渠道'][0]))  # 将docx文件写入zip文件
            os.remove(u'{0}_{1}.docx'.format(project_name, replace_content_dict['渠道'][0]))  # 删除文件
            z.close()
            f = open('authorizations.zip')
            data = f.read()  # 读取压缩文件的数据
            f.close()
            return request.make_response(
                [data, data],
                headers=[
                    ('Content-Disposition', 'attachment; filename="%s"'
                     % 'authorization.zip'),
                    ('Content-Type', self.content_type)
                ],
                cookies={'fileToken': token}
            )


    @http.route('/dtdream_authorization/dtdream_shouhou_export', type='http', auth='public')
    def dtdream_authorization_export(self, data, token):
        active_ids = data.split(',')
        dtdream_authorization_apply = request.env['dtdream.project.bid.authorize.apply'].search(
            [('id', '=', int(active_ids[0]))])
        from docx import Document
        import sys
        reload(sys)
        sys.setdefaultencoding('utf-8')
        auth = dtdream_authorization_apply
        project_name = auth.bidding_rep_pro_name
        bidding_number = auth.bidding_number
        date = auth.bidding_time.split(u'-')
        yearnum = str(auth.yearnum)
        doc = Document('myaddons/dtdream_project_bid_authorize_apply/static/file/shouhou.docx')
        paragraphs = doc.paragraphs
        replace_content = ['招标项目名称', '项目名称', 'bidnum', '2016年6月12日', 'year', 'month', 'day','ynu']
        replace_content_dict = {'招标项目名称': project_name, '项目名称': project_name, 'bidnum': bidding_number,'ynu':yearnum, 'year': date[0],
                                'month': date[1], 'day': date[2]}
        fp = StringIO()
        for p in paragraphs:
            for run in p.runs:
                for rc in replace_content:
                    if rc in run.text:
                        run.text = run.text.replace(rc, replace_content_dict[rc])
        doc.save(fp)
        fp.seek(0)
        data = fp.read()
        fp.close()
        return request.make_response(
            data,
            headers=[
                ('Content-Disposition', 'attachment; filename="shouhou.docx"'
                 ),
                ('Content-Type', self.content_type)
            ],
            cookies={'fileToken': token}
        )